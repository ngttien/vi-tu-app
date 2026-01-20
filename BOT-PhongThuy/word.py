from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_placeholder_image(doc, label):
    """Hàm tạo khung trống để chèn ảnh sau này"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n[... KHOẢNG TRỐNG CHÈN ẢNH: {label} ...]\n(Vui lòng xóa dòng này và paste ảnh minh họa vào đây)\n")
    run.font.color.rgb = RGBColor(255, 0, 0) # Màu đỏ để dễ nhìn
    run.font.bold = True
    run.font.size = Pt(11)
    # Thêm khoảng cách dòng để dễ paste
    p_spacing = doc.add_paragraph() 
    p_spacing.paragraph_format.space_after = Pt(20)

def create_tech_spec_doc():
    doc = Document()

    # --- 1. TIÊU ĐỀ ---
    title = doc.add_heading('TÀI LIỆU KỸ THUẬT & DEMO\nHỆ THỐNG BOT PHONG THỦY TỰ ĐỘNG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Phiên bản: 1.0')
    doc.add_paragraph('Mô tả: Hệ thống tự động hóa quy trình tư vấn Phong thủy/Tử vi sử dụng AI thông qua Browser Automation.')
    doc.add_page_break()

    # --- 2. GIAO DIỆN NGƯỜI DÙNG (UI) ---
    doc.add_heading('1. Giao diện Người dùng (UI)', level=1)
    doc.add_paragraph('Mô tả: Giao diện dòng lệnh (CLI) hoặc giao diện nhập liệu nơi người dùng cung cấp thông tin ban đầu (Họ tên, Ngày sinh, Giờ sinh...).')
    
    doc.add_heading('1.1. Màn hình nhập liệu / Menu chính', level=2)
    doc.add_paragraph('Hệ thống khởi chạy với Menu tùy chọn các luồng xử lý:')
    add_placeholder_image(doc, "GIAO DIỆN MENU CMD / MAIN.PY")

    # --- 3. LUỒNG 1: XỬ LÝ ĐẦU VÀO & GOOGLE SHEET ---
    doc.add_heading('2. Luồng xử lý 1: Tiếp nhận & Lưu trữ', level=1)
    
    doc.add_heading('2.1. Nhận File JSON đầu vào', level=2)
    doc.add_paragraph('Hệ thống nhận dữ liệu thô từ khách hàng, tạo UUID định danh và lưu thành file JSON tại thư mục local data/raw.')
    add_placeholder_image(doc, "ẢNH FILE JSON ĐƯỢC TẠO TRONG FOLDER DATA/RAW")

    doc.add_heading('2.2. Cập nhật Google Sheet (Database)', level=2)
    doc.add_paragraph('Thông tin cơ bản (Tên, Email, Trạng thái PENDING) được đẩy lên Google Sheet theo thời gian thực để quản lý hàng đợi.')
    add_placeholder_image(doc, "ẢNH GOOGLE SHEET VỚI TRẠNG THÁI PENDING")

    # --- 4. LUỒNG 2: AI PROCESSING (BROWSER AUTOMATION) ---
    doc.add_heading('3. Luồng xử lý 2: AI Processing (Browser Automation)', level=1)
    doc.add_paragraph('Sử dụng Playwright để điều khiển trình duyệt Chrome thật, kết nối với Gemini/ChatGPT mà không cần API Key.')

    doc.add_heading('3.1. Kịch bản Prompt (Context Injection)', level=2)
    doc.add_paragraph('Hệ thống tự động đọc file JSON và file kiến thức (nếu có), ghép thành một Prompt hoàn chỉnh yêu cầu AI trả về định dạng JSON.')
    add_placeholder_image(doc, "ẢNH CODE PYTHON PHẦN PROMPT TRONG GEMINI_WORKER.PY")

    doc.add_heading('3.2. Quá trình gửi và nhận phản hồi trên Browser', level=2)
    doc.add_paragraph('Bot tự động mở trình duyệt, paste nội dung vào khung chat và đợi câu trả lời.')
    add_placeholder_image(doc, "ẢNH CHỤP MÀN HÌNH TRÌNH DUYỆT ĐANG CHAT VỚI GEMINI/GPT")

    doc.add_heading('3.3. Kết quả JSON trả về từ AI', level=2)
    doc.add_paragraph('Kết quả thô từ AI được làm sạch, validate và lưu thành file JSON kết quả (Processed Data).')
    add_placeholder_image(doc, "ẢNH FILE JSON KẾT QUẢ TRONG FOLDER DATA/PROCESSED")

    # --- 5. LUỒNG 3: TẠO PDF & GỬI EMAIL ---
    doc.add_heading('4. Luồng xử lý 3: Xuất bản & Bàn giao', level=1)
    doc.add_paragraph('Đây là bước cuối cùng: Biến dữ liệu JSON thành văn bản người đọc hiểu được và gửi đi.')

    doc.add_heading('4.1. Template Word/PDF mẫu', level=2)
    doc.add_paragraph('Hệ thống sử dụng một file Word mẫu (.docx) có chứa các biến {{variable}} tương ứng với các trường trong file JSON.')
    add_placeholder_image(doc, "ẢNH FILE WORD TEMPLATE (CÓ CÁC BIẾN {{...}})")

    doc.add_heading('4.2. File PDF hoàn chỉnh', level=2)
    doc.add_paragraph('Sau khi điền dữ liệu (Fill data), hệ thống convert file sang PDF.')
    add_placeholder_image(doc, "ẢNH FILE PDF KẾT QUẢ ĐẸP MẮT")

    doc.add_heading('4.3. Gửi Email tự động', level=2)
    doc.add_paragraph('Hệ thống login SMTP Server, đính kèm file PDF và gửi lời chào tới khách hàng. Cập nhật trạng thái trên Sheet thành DONE.')
    add_placeholder_image(doc, "ẢNH HỘP THƯ EMAIL KHÁCH HÀNG NHẬN ĐƯỢC FILE")

    # --- 6. TỔNG KẾT ---
    doc.add_heading('5. Tổng kết & Mã nguồn', level=1)
    doc.add_paragraph('Toàn bộ hệ thống hoạt động khép kín, không phụ thuộc API trả phí, dễ dàng mở rộng.')
    add_placeholder_image(doc, "ẢNH CẤU TRÚC THƯ MỤC PROJECT (TREE FOLDER)")

    # Lưu file
    file_name = 'Tai_lieu_Demo_Bot_Phong_Thuy.docx'
    doc.save(file_name)
    print(f"[OK] Đã tạo xong file: {file_name}")

if __name__ == "__main__":
    create_tech_spec_doc()